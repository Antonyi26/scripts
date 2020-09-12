# программа распределения работников
import os
import re
from random import shuffle
from docx import Document

######################################################

class Table:
    def __init__(self):
        self.employers = []
        self.rows = []
        self.dateMapped = {}
        self.regionSet = set()
        self.SPbSet = set()
        
    def add(self, tableRow):
        self.rows.append(tableRow)
        ind = len(self.rows) - 1
        # дата
        val = tableRow.date.toString()
        if val not in self.dateMapped:
            self.dateMapped[val] = []
        self.dateMapped[val].append(ind)
        # город или город + улица, если это Питер
        val = tableRow.adress.getCity()
        if val == 'Санкт-Петербург':
            val += ' --> ' + tableRow.adress.getStreet()
            self.SPbSet.add(val)
        else:
            self.regionSet.add(val)
        #if val not in self.cityMapped:
        #    self.cityMapped[val] = []
        #self.cityMapped[val].append(ind)
    
    def addEmployers(self, employers):
        self.employers = employers
    
    def arrangeEmploers(self):
        for date in self.dateMapped:
            # распределяем работников по адресам
            empl = self.employers.copy()
            shuffle(empl)
            emplOnCityMap = {}
            # сначала пригоро
            i = 0
            for val in self.regionSet:
                emplOnCityMap[val] = empl[i % len(empl)]
                i += 1
            # затем СПб
            for val in self.SPbSet:
                emplOnCityMap[val] = empl[i % len(empl)]
                i += 1
            # заносим в локальную таблицу
            for ind in self.dateMapped[date]:
                val = self.rows[ind].adress.getCity()
                if val == 'Санкт-Петербург':
                    val += ' --> ' + self.rows[ind].adress.getStreet()
                self.rows[ind].employer = emplOnCityMap[val]  
    
    def show(self):
        for row in self.rows:
            row.show()
            print()
        
    def rowsCount(self):
        return len(self.rows)
        
    def cell(self, row, col):
        return self.rows[row].data(col)
        
 ######################################################

class TableRow:
    def __init__(self, id, date, adress, fiderNum, employer=None):
        self.id = int(id)
        self.date = Date(date)
        self.adress = Adress(adress)
        self.fiderNum = str(fiderNum)
        self.employer = employer
        
    def data(self, col):
        if col == 0: return self.id
        if col == 1: return self.date.toString()
        if col == 2: return self.adress.getRawAdress()
        if col == 3: return self.fiderNum
        if col == 4: return self.employer
        return None
        
    def colCount(self):
        return 5
        
    def show(self):
        print(self.id)
        self.date.show()
        print(self.fiderNum)
        self.adress.show()


######################################################

class Date:
    def __init__(self, date):
        date = re.search(r'(\d+)\.(\d+)\.(\d+)', date, flags=0)
        if date:
            self.day = date.group(1)
            if len(self.day) == 1: 
                self.day = '0' + self.day 
            self.month = date.group(2)
            if len(self.month) == 1: 
                self.month = '0' + self.month         
            self.year = date.group(3)
            if len(self.year) == 2: 
                self.year = '20' + self.year
    
    def show(self):
        print(self.day, self.month, self.year)
        
    def toString(self):
        return self.day + self.month + self.year

######################################################

class Adress:
    def __init__(self, adress):
        self.rawAdress = adress
    
        res = re.search(r'[Гг]\.\s*(\w+)\s*,', adress, flags=0)
        self.city = res.group(1) if res else None
        if self.city is None:
            self.city = 'Санкт-Петербург'
        
        res = re.search(r'[Уу]л\.\s*([\w\s-]+)\s*,', adress, flags=0)
        if not res:
            res = re.search(r'[Пп]р\.\s*([\w\s-]+)\s*,', adress, flags=0)
        self.street = res.group(1) if res else None
        
        res = re.search(r'[Дд]\.\s*(\d+)', adress, flags=0)
        self.build = res.group(1) if res else None

    def show(self):
        print(self.city, self.street, self.build)
        
    def getCity(self):
        return self.city
       
    def getStreet(self):
        return self.street
      
    def getRawAdress(self):
        return self.rawAdress

######################################################


class Employer:
    def __init__(self, employer):
        employer = re.search(r'(\w+)', employer, flags=0)
        self.sorname = str(employer.group(1))

    def show(self):
        print(self.sorname)

######################################################

# делаем текущую папку рабочей
os.chdir(os.path.abspath(os.path.dirname(__file__)))
# читаем документ
wordDoc = Document('test.docx')
rows = len(wordDoc.tables[0].rows)
table = Table()

for row in range(1, rows):
    id = wordDoc.tables[0].rows[row].cells[0].text
    date = wordDoc.tables[0].rows[row].cells[1].text
    adress = wordDoc.tables[0].rows[row].cells[2].text
    fiderNum = wordDoc.tables[0].rows[row].cells[3].text
    table.add( TableRow(id, date, adress, fiderNum) )

# распределяе работников
with open('employers.txt', 'r', encoding="utf-8") as f:
    table.addEmployers( f.readlines() )
table.arrangeEmploers()

# сохраняем в новый документ
newDoc = Document()
newDocTable = newDoc.add_table(cols = 5, rows = table.rowsCount() + 1)
newDocTable.style = 'Table Grid'  # применяем стиль для таблицы
cols = len(newDocTable.columns)
rows = len(newDocTable.rows)
# нулевая строка - заголовок таблицы
for col in range(cols-1):
    newDocTable.cell(0, col).text = wordDoc.tables[0].rows[0].cells[col].text
newDocTable.cell(0, cols-1).text = 'Инженеры'
# заполняем тело таблицы
for row in range(1, rows):
    for col in range(cols):
        newDocTable.cell(row, col).text = str( table.cell(row-1, col) )
    
newDoc.save('result.docx')
#for row in range(1, table.rowsCount())



